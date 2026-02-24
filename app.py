import streamlit as st
import PyPDF2
import docx
import time  # <-- THÊM THƯ VIỆN ĐO THỜI GIAN
from io import BytesIO
from summarizer_ai import TextSummarizer
from textrank_summarizer import TextRankSummarizer
from text_cleaner import TextPreprocessor

# ==========================================
# 1. CẤU HÌNH TRANG VÀ GIAO DIỆN
# ==========================================
st.set_page_config(page_title="AI Summarizer Pro", page_icon="📝", layout="wide")

@st.cache_resource
def load_models():
    return TextSummarizer(), TextRankSummarizer(), TextPreprocessor()

ai_summarizer, textrank_summarizer, text_cleaner = load_models()

# ==========================================
# 2. HÀM XỬ LÝ TRÍCH XUẤT VĂN BẢN
# ==========================================
def extract_text_from_file(uploaded_file):
    """Đọc nội dung từ file TXT, PDF hoặc DOCX"""
    try:
        filename = uploaded_file.name
        if filename.endswith('.txt'):
            return uploaded_file.getvalue().decode("utf-8")
        
        elif filename.endswith('.pdf'):
            pdf_reader = PyPDF2.PdfReader(BytesIO(uploaded_file.read()))
            text = ""
            for page in pdf_reader.pages:
                content = page.extract_text()
                if content:
                    text += content + "\n"
            return text
        
        elif filename.endswith('.docx'):
            doc = docx.Document(BytesIO(uploaded_file.read()))
            return "\n".join([para.text for para in doc.paragraphs])
    except Exception as e:
        st.error(f"Lỗi khi đọc file: {e}")
        return ""
    return ""

# ==========================================
# 3. GIAO DIỆN NGƯỜI DÙNG (UI)
# ==========================================
st.title("📝 Hệ thống Tóm tắt Văn bản Thông minh")
st.markdown("Hệ thống hỗ trợ tóm tắt đa định dạng, cho phép tùy chỉnh độ dài văn bản theo nhu cầu người đọc.")

# --- THANH ĐIỀU KHIỂN BÊN TRÁI (SIDEBAR) ---
st.sidebar.header("⚙️ Cấu hình tóm tắt")
summary_length = st.sidebar.slider("Độ dài tóm tắt mong muốn (số từ):", 30, 1000, 100, help="AI sẽ cố gắng tóm tắt sát với số lượng từ này nhất.")
method = st.sidebar.selectbox(
    "Chọn phương thức tóm tắt:", 
    ["Thông minh (AI T5 - Viết lại câu)", "Trích xuất ý chính (TextRank - Giữ nguyên câu)"]
)

st.sidebar.markdown("---")
st.sidebar.info("""
**Hướng dẫn:**
1. Tải file tài liệu hoặc dán văn bản.
2. Chọn độ dài và phương thức.
3. Nhấn nút 'Tiến hành Tóm tắt'.
""")

# --- KHU VỰC NHẬP DỮ LIỆU ---
st.subheader("📥 Dữ liệu đầu vào")
uploaded_file = st.file_uploader("📂 Tải lên tài liệu (Hỗ trợ: PDF, DOCX, TXT)", type=["pdf", "docx", "txt"])

input_content = ""
if uploaded_file is not None:
    with st.spinner("Đang trích xuất dữ liệu từ file..."):
        input_content = extract_text_from_file(uploaded_file)
        if input_content:
            st.success(f"✅ Đã nhận diện nội dung từ file: {uploaded_file.name}")

input_text = st.text_area(
    "Nội dung văn bản cần xử lý:", 
    value=input_content, 
    height=300, 
    placeholder="Nhập hoặc dán văn bản của bạn tại đây..."
)

# --- NÚT BẤM VÀ LOGIC XỬ LÝ ---
col1, col2 = st.columns([1, 4])
with col1:
    btn_run = st.button("🚀 Tiến hành Tóm tắt", type="primary")

if btn_run:
    if len(input_text.strip()) < 50:
        st.warning("⚠️ Văn bản quá ngắn (dưới 50 ký tự) để thực hiện tóm tắt chất lượng.")
    else:
        with st.spinner("🤖 AI đang đọc và phân tích văn bản..."):
            
            # --- BẮT ĐẦU ĐO THỜI GIAN ---
            start_time = time.time()
            
            cleaned_text = text_cleaner.clean_text(input_text)
            
            if method == "Thông minh (AI T5 - Viết lại câu)":
                result = ai_summarizer.summarize(cleaned_text, max_len=summary_length)
            else:
                num_sentences = max(1, summary_length // 20)
                result = textrank_summarizer.summarize(cleaned_text, num_sentences=num_sentences)
            
            # --- KẾT THÚC ĐO THỜI GIAN ---
            end_time = time.time()
            processing_time = round(end_time - start_time, 2)
            
            # TÍNH TOÁN TỶ LỆ NÉN
            original_word_count = len(cleaned_text.split())
            summary_word_count = len(result.split())
            if original_word_count > 0:
                compression_ratio = round((summary_word_count / original_word_count) * 100, 1)
            else:
                compression_ratio = 0
            
            # --- HIỂN THỊ KẾT QUẢ ---
            st.markdown("---")
            st.subheader("📄 Kết quả tóm tắt:")
            st.success(result)
            
            keywords = textrank_summarizer.extract_keywords(cleaned_text, num_keywords=5)
            if keywords:
                tags_html = " ".join([f"`#{kw.capitalize()}`" for kw in keywords])
                st.markdown(f"**🔑 Từ khóa chính:** {tags_html}")
            
            # ==========================================
            # BẢNG THÔNG SỐ SO SÁNH (MỚI)
            # ==========================================
            st.markdown("### 📊 Thông số hiệu năng")
            metric_col1, metric_col2, metric_col3 = st.columns(3)
            metric_col1.metric(label="⏱️ Thời gian xử lý", value=f"{processing_time} giây")
            metric_col2.metric(label="📉 Tỷ lệ nén", value=f"{compression_ratio}%")
            metric_col3.metric(label="📝 Độ dài (Tóm tắt / Gốc)", value=f"{summary_word_count} / {original_word_count} từ")

            # --- TÍNH NĂNG XUẤT FILE ---
            st.markdown("### 📥 Xuất kết quả")
            col_txt, col_word, _ = st.columns([1, 1, 2])
            
            with col_txt:
                st.download_button(
                    label="📄 Tải file Text (.txt)",
                    data=result,
                    file_name="Ket_qua_tom_tat.txt",
                    mime="text/plain",
                    use_container_width=True
                )
            with col_word:
                doc_result = docx.Document()
                doc_result.add_heading('Bản Tóm Tắt Tự Động (AI Summarizer)', level=1)
                doc_result.add_paragraph(result)
                bio = BytesIO()
                doc_result.save(bio)
                bio.seek(0)
                st.download_button(
                    label="📘 Tải file Word (.docx)",
                    data=bio,
                    file_name="Ket_qua_tom_tat.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True
                )

st.markdown("---")
st.caption("Hệ thống tóm tắt văn bản tự động - Nghiên cứu so sánh AI")